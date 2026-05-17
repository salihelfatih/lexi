"""API and end-to-end verification tests."""

import os
import tempfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4

import pytest
import httpx
from docx import Document as DocxDocument
from sqlalchemy.exc import OperationalError

TEST_DB_PATH = Path(tempfile.gettempdir()) / f"lexi_api_tests_{os.getpid()}.sqlite"
if "TEST_DATABASE_URL" in os.environ:
    os.environ["DATABASE_URL"] = os.environ["TEST_DATABASE_URL"]
else:
    os.environ["DATABASE_URL"] = f"sqlite:///{TEST_DB_PATH}"
os.environ["LLM_PROVIDER"] = "fake"
os.environ["LLM_MODEL_NAME"] = "fake-grounded-summary-v1"
os.environ["RAG_VECTOR_BACKEND"] = "in_memory"
os.environ["RAG_EMBEDDING_BACKEND"] = "deterministic"

from backend.database import Base, SessionLocal, engine  # noqa: E402
from backend.main import app  # noqa: E402
from backend.models import (  # noqa: E402
    Clause,
    ConsentRecord,
    Document,
    DocumentMetadata,
    DocumentSummary,
    JobStatus,
    RiskSignal,
    User,
)  # noqa: E402
from backend.security import get_current_user  # noqa: E402


async def _test_user():
    return SimpleNamespace(id=uuid4(), email="tester@example.com", is_active=True)


@pytest.fixture(scope="module", autouse=True)
def database_schema():
    """Ensure the API has tables available for integration tests."""
    try:
        Base.metadata.drop_all(bind=engine)
        Base.metadata.create_all(bind=engine)
    except OperationalError as exc:
        pytest.skip(f"Database is not available for API integration tests: {exc}")
    yield
    engine.dispose()
    if "TEST_DATABASE_URL" not in os.environ:
        TEST_DB_PATH.unlink(missing_ok=True)


def _api_client():
    transport = httpx.ASGITransport(app=app)
    return httpx.AsyncClient(transport=transport, base_url="http://testserver")


@pytest.fixture
def authenticated_uploads():
    app.dependency_overrides[get_current_user] = _test_user
    yield
    app.dependency_overrides.pop(get_current_user, None)


@pytest.fixture
def eager_worker():
    from backend.celery_app import celery_app

    previous_always_eager = celery_app.conf.task_always_eager
    previous_eager_propagates = celery_app.conf.task_eager_propagates

    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True

    yield

    celery_app.conf.task_always_eager = previous_always_eager
    celery_app.conf.task_eager_propagates = previous_eager_propagates


def _sample_lease_docx() -> BytesIO:
    doc = DocxDocument()
    doc.add_heading("Ontario Standard Lease", level=1)
    paragraphs = [
        "This residential tenancy agreement is governed by the Residential Tenancies Act, Ontario, and may be reviewed by the Landlord and Tenant Board.",
        "Landlord: Alex Smith",
        "Tenant: Jane Doe",
        "Rental unit: 123 King Street, Toronto, ON",
        "Lease term start date: 01/01/2026",
        "Lease term end date: 12/31/2026",
        "Monthly rent: $2450 per month",
        "1. Rent payment. The tenant must pay monthly rent on the first day of each month.",
        "2. Maintenance and repair. The landlord is responsible for maintenance and repair of the rental unit.",
        "3. Entry and access. The landlord may enter for inspection only with lawful notice.",
        "4. Utilities. Hydro, water, gas, and electricity responsibilities are listed in this lease.",
        "5. Notice to end. Either party must follow Ontario notice requirements to terminate the lease.",
        "6. Subletting. The tenant may request assignment or sublet approval in writing.",
    ]

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    content = BytesIO()
    doc.save(content)
    content.seek(0)
    return content


def _random_unsupported_text() -> list[str]:
    return [
        "Weekly bakery production report for the downtown kitchen team.",
        "Croissant batches were counted before opening, muffin trays were labeled, and delivery crates were checked against supplier invoices.",
        "The afternoon notes cover oven cleaning, refrigerator temperature logs, packaging counts, staff training reminders, ingredient substitutions, and customer feedback.",
        "Community event orders include fruit tarts, sandwich platters, coffee service, seasonal menu experiments, and a checklist for tomorrow morning.",
        "The closing summary records register reconciliation, broken tray replacement, pantry restocking, and schedule notes for the next production cycle.",
    ]


def _random_unsupported_docx() -> BytesIO:
    doc = DocxDocument()
    doc.add_heading("Bakery Operations Packet", level=1)
    for paragraph in _random_unsupported_text():
        doc.add_paragraph(paragraph)

    content = BytesIO()
    doc.save(content)
    content.seek(0)
    return content


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _random_unsupported_pdf() -> BytesIO:
    lines = ["Bakery Operations Packet", *_random_unsupported_text()]
    text_stream = "\n".join(
        [
            "BT",
            "/F1 10 Tf",
            "50 760 Td",
            "14 TL",
            *[item for line in lines for item in (f"({_escape_pdf_text(line)}) Tj", "T*")],
            "ET",
        ]
    )
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(text_stream.encode('utf-8'))} >>\nstream\n{text_stream}\nendstream",
    ]

    pdf = "%PDF-1.4\n"
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf.encode("utf-8")))
        pdf += f"{index} 0 obj\n{obj}\nendobj\n"

    xref_offset = len(pdf.encode("utf-8"))
    pdf += f"xref\n0 {len(objects) + 1}\n"
    pdf += "0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n"
    pdf += f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n"
    pdf += f"startxref\n{xref_offset}\n%%EOF\n"

    return BytesIO(pdf.encode("utf-8"))


def _cleanup_e2e_records(email, document_id=None):
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.email == email).first()
        if document_id:
            db.query(RiskSignal).filter(RiskSignal.document_id == document_id).delete()
            db.query(Clause).filter(Clause.document_id == document_id).delete()
            db.query(DocumentSummary).filter(DocumentSummary.document_id == document_id).delete()
            db.query(DocumentMetadata).filter(DocumentMetadata.document_id == document_id).delete()
            db.query(JobStatus).filter(JobStatus.document_id == document_id).delete()
            db.query(ConsentRecord).filter(ConsentRecord.document_id == document_id).delete()
            db.query(Document).filter(Document.id == document_id).delete()

        if user:
            db.delete(user)

        db.commit()
    finally:
        db.close()


def _assert_risk_language_stays_informational(risk_sense):
    generated_text = " ".join(
        [
            risk_sense["top_risks_summary"],
            *[
                f"{risk['title']} {risk['reason']}"
                for risk in risk_sense["risks"]
            ],
        ]
    ).lower()
    forbidden_phrases = [
        "illegal",
        "unlawful",
        "unenforceable",
        "against the law",
        "legal advice",
        "will win",
        "will lose",
        "will be evicted",
        "guarantee",
        "must sign",
        "do not sign",
        "outcome",
    ]

    for phrase in forbidden_phrases:
        assert phrase not in generated_text


@pytest.mark.asyncio
async def test_health_check():
    """Test health check endpoint."""
    async with _api_client() as client:
        response = await client.get("/health")

    assert response.status_code == 200
    assert response.json() == {"status": "healthy", "service": "lexi"}


@pytest.mark.asyncio
async def test_upload_invalid_file_size(authenticated_uploads, monkeypatch):
    """Test upload with oversized file."""
    from backend.api.v1 import documents

    monkeypatch.setattr(documents.settings, "max_file_size_mb", 0)

    async with _api_client() as client:
        response = await client.post(
            "/api/v1/documents/upload",
            files={"file": ("large.pdf", b"%PDF-1.4\n", "application/pdf")},
        )

    assert response.status_code == 413


@pytest.mark.asyncio
async def test_upload_invalid_format(authenticated_uploads):
    """Test upload with unsupported format."""
    async with _api_client() as client:
        response = await client.post(
            "/api/v1/documents/upload", files={"file": ("test.txt", b"test content", "text/plain")}
        )

    assert response.status_code == 415


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("filename", "content_factory", "mime_type", "expected_format"),
    [
        (
            "bakery-operations.docx",
            _random_unsupported_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        ),
        ("bakery-operations.pdf", _random_unsupported_pdf, "application/pdf", "pdf"),
    ],
)
async def test_supported_format_random_content_completes_unsupported(
    eager_worker,
    filename,
    content_factory,
    mime_type,
    expected_format,
):
    """Random content in supported containers completes as unsupported by Lexi."""
    email = f"unsupported-e2e-{uuid4().hex}@example.com"
    password = "unsupported-proof-password"
    document_id = None

    try:
        async with _api_client() as client:
            register_response = await client.post(
                "/api/v1/auth/register",
                json={"email": email, "password": password},
            )
            assert register_response.status_code == 201

            login_response = await client.post(
                "/api/v1/auth/login",
                data={"username": email, "password": password},
            )
            assert login_response.status_code == 200
            auth_headers = {"Authorization": f"Bearer {login_response.json()['access_token']}"}

            upload_response = await client.post(
                "/api/v1/documents/upload",
                headers=auth_headers,
                files={"file": (filename, content_factory(), mime_type)},
            )
            assert upload_response.status_code == 201
            upload_payload = upload_response.json()
            document_id = upload_payload["document_id"]
            assert upload_payload["format"] == expected_format

            consent_response = await client.post(
                f"/api/v1/documents/{document_id}/consent",
                headers=auth_headers,
                json={"processing_consent": True, "storage_consent": False},
            )
            assert consent_response.status_code == 200

            status_response = await client.get(
                f"/api/v1/jobs/{document_id}/status",
                headers=auth_headers,
            )
            assert status_response.status_code == 200
            assert status_response.json()["status"] == "complete"

            results_response = await client.get(
                f"/api/v1/documents/{document_id}/results",
                headers=auth_headers,
            )
            assert results_response.status_code == 200
            results = results_response.json()
            assert results["document_type"] == "unknown"
            assert results["classification_confidence"] >= 90
            assert results["summary"] is None
            assert results["risk_sense"] is None
            assert results["clauses"] == []
            assert results["total_clauses"] == 0

            ask_response = await client.post(
                f"/api/v1/documents/{document_id}/ask",
                headers=auth_headers,
                json={"question": "What does this document say about rent?"},
            )
            assert ask_response.status_code == 400
            assert "unsupported by Lexi" in ask_response.json()["detail"]

            documents_response = await client.get(
                "/api/v1/documents",
                headers=auth_headers,
            )
            assert documents_response.status_code == 200
            documents_payload = documents_response.json()["documents"]
            assert len(documents_payload) == 1
            assert documents_payload[0]["document_type"] == "unknown"
            assert documents_payload[0]["job_status"] == "complete"
            assert documents_payload[0]["total_clauses"] == 0
            assert "extracted_text" not in documents_payload[0]
            assert "clauses" not in documents_payload[0]

            db = SessionLocal()
            try:
                assert (
                    db.query(DocumentSummary)
                    .filter(DocumentSummary.document_id == document_id)
                    .first()
                    is None
                )
                assert db.query(Clause).filter(Clause.document_id == document_id).count() == 0
                assert (
                    db.query(DocumentMetadata)
                    .filter(DocumentMetadata.document_id == document_id)
                    .first()
                    is None
                )
                assert (
                    db.query(RiskSignal).filter(RiskSignal.document_id == document_id).count()
                    == 0
                )
            finally:
                db.close()

            delete_response = await client.delete(
                f"/api/v1/documents/{document_id}",
                headers=auth_headers,
            )
            assert delete_response.status_code == 204
    finally:
        _cleanup_e2e_records(email, document_id)


@pytest.mark.asyncio
async def test_end_to_end_lease_verification_path(eager_worker):
    """Prove register/login, upload, consent, worker, results, and cleanup in one path."""
    email = f"lease-e2e-{uuid4().hex}@example.com"
    password = "lease-proof-password"
    document_id = None

    try:
        async with _api_client() as client:
            register_response = await client.post(
                "/api/v1/auth/register",
                json={"email": email, "password": password},
            )
            assert register_response.status_code == 201
            assert register_response.json()["email"] == email

            login_response = await client.post(
                "/api/v1/auth/login",
                data={"username": email, "password": password},
            )
            assert login_response.status_code == 200
            token = login_response.json()["access_token"]
            auth_headers = {"Authorization": f"Bearer {token}"}

            empty_documents_response = await client.get(
                "/api/v1/documents",
                headers=auth_headers,
            )
            assert empty_documents_response.status_code == 200
            assert empty_documents_response.json() == {"documents": []}

            upload_response = await client.post(
                "/api/v1/documents/upload",
                headers=auth_headers,
                files={
                    "file": (
                        "sample-ontario-lease.docx",
                        _sample_lease_docx(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            assert upload_response.status_code == 201
            upload_payload = upload_response.json()
            document_id = upload_payload["document_id"]
            assert upload_payload["format"] == "docx"

            uploaded_documents_response = await client.get(
                "/api/v1/documents",
                headers=auth_headers,
            )
            assert uploaded_documents_response.status_code == 200
            uploaded_documents = uploaded_documents_response.json()["documents"]
            assert len(uploaded_documents) == 1
            assert uploaded_documents[0]["document_id"] == document_id
            assert uploaded_documents[0]["filename"] == "sample-ontario-lease.docx"
            assert uploaded_documents[0]["job_status"] in {None, "pending"}
            assert "extracted_text" not in uploaded_documents[0]
            assert "clauses" not in uploaded_documents[0]

            temp_path = Path(tempfile.gettempdir()) / "lexi_uploads" / document_id
            assert temp_path.exists()

            consent_response = await client.post(
                f"/api/v1/documents/{document_id}/consent",
                headers=auth_headers,
                json={"processing_consent": True, "storage_consent": False},
            )
            assert consent_response.status_code == 200
            assert consent_response.json()["storage_mode"] == "ephemeral"

            status_response = await client.get(
                f"/api/v1/jobs/{document_id}/status",
                headers=auth_headers,
            )
            assert status_response.status_code == 200
            assert status_response.json()["status"] == "complete"

            results_response = await client.get(
                f"/api/v1/documents/{document_id}/results",
                headers=auth_headers,
            )
            assert results_response.status_code == 200
            results = results_response.json()
            assert results["document_type"] == "ontario_residential_lease"
            assert results["classification_confidence"] >= 70
            assert results["metadata"]["tenant_names"] == ["Jane Doe"]
            assert results["metadata"]["landlord_names"] == ["Alex Smith"]
            assert results["metadata"]["property_address"] == "123 King Street, Toronto, ON"
            assert results["metadata"]["monthly_rent"] == "$2450"
            assert results["metadata"]["lease_start_date"] == "2026-01-01"
            assert results["metadata"]["lease_end_date"] == "2026-12-31"
            assert results["total_clauses"] >= 6
            assert any(clause["clause_type"] == "maintenance" for clause in results["clauses"])
            assert any(clause["clause_type"] == "access" for clause in results["clauses"])
            assert results["summary"] == {
                "text": (
                    "This ontario residential lease names Jane Doe as tenant and Alex Smith as "
                    "landlord; identifies the rental unit as 123 King Street, Toronto, ON; "
                    "lists monthly rent of $2450; sets a lease term from 01/01/2026 to "
                    "12/31/2026; includes clauses about rent payment, maintenance, entry and "
                    "access, utilities, ending the lease, and subletting."
                ),
                "provider": "fake",
                "model": "fake-grounded-summary-v1",
                "grounded_in": "clauses",
                "source_count": results["total_clauses"],
            }
            risk_sense = results["risk_sense"]
            assert risk_sense is not None
            assert risk_sense["top_risks_summary"].startswith("RiskSense found")
            assert risk_sense["confidence_rollup"]["extraction"] > 0
            assert risk_sense["confidence_rollup"]["classification"] >= 70
            assert risk_sense["confidence_rollup"]["metadata"] > 0
            assert risk_sense["confidence_rollup"]["risk_signals"] > 0
            assert risk_sense["confidence_rollup"]["overall"] > 0
            assert risk_sense["risks"]
            assert {risk["severity"] for risk in risk_sense["risks"]}.issubset(
                {"low", "medium", "high"}
            )
            assert any(risk["rule_id"] == "landlord_access" for risk in risk_sense["risks"])
            clause_texts = {clause["clause_text"] for clause in results["clauses"]}
            for risk in risk_sense["risks"]:
                assert risk["confidence"] > 0
                assert risk["reason"]
                assert risk["source_clause"]["clause_text"] in clause_texts
            _assert_risk_language_stays_informational(risk_sense)

            documents_response = await client.get(
                "/api/v1/documents",
                headers=auth_headers,
            )
            assert documents_response.status_code == 200
            documents_payload = documents_response.json()["documents"]
            assert len(documents_payload) == 1
            assert documents_payload[0]["job_status"] == "complete"
            assert documents_payload[0]["document_type"] == "ontario_residential_lease"
            assert documents_payload[0]["has_results"] is True
            assert documents_payload[0]["total_clauses"] >= 6
            assert "extracted_text" not in documents_payload[0]

            ask_response = await client.post(
                f"/api/v1/documents/{document_id}/ask",
                headers=auth_headers,
                json={
                    "question": "What is the monthly rent?",
                    "user_context": "My landlord told me rent might be $3000.",
                },
            )
            assert ask_response.status_code == 200
            answer = ask_response.json()
            assert answer["is_answered"] is True
            assert "$2450" in answer["answer"]
            assert "$3000" not in answer["answer"]
            assert answer["citations"]
            assert "$2450" in answer["citations"][0]["text"]
            assert "kept separate" in answer["user_context_note"]

            missing_answer_response = await client.post(
                f"/api/v1/documents/{document_id}/ask",
                headers=auth_headers,
                json={"question": "Does this document say parking is included?"},
            )
            assert missing_answer_response.status_code == 200
            missing_answer = missing_answer_response.json()
            assert missing_answer["is_answered"] is False
            assert "I don't know from this document" in missing_answer["answer"]
            assert missing_answer["citations"] == []

            delete_response = await client.delete(
                f"/api/v1/documents/{document_id}",
                headers=auth_headers,
            )
            assert delete_response.status_code == 204
            assert not temp_path.exists()

            deleted_results_response = await client.get(
                f"/api/v1/documents/{document_id}/results",
                headers=auth_headers,
            )
            assert deleted_results_response.status_code == 404

            deleted_job_response = await client.get(
                f"/api/v1/jobs/{document_id}/status",
                headers=auth_headers,
            )
            assert deleted_job_response.status_code == 404

            deleted_documents_response = await client.get(
                "/api/v1/documents",
                headers=auth_headers,
            )
            assert deleted_documents_response.status_code == 200
            assert deleted_documents_response.json() == {"documents": []}
    finally:
        _cleanup_e2e_records(email, document_id)
